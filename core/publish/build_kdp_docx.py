# © 2026 Sabino Gervasio · os.getenv("GMAIL_USER", "your@email.com") · ${GUMROAD_USERNAME}.gumroad.com
# Data: 2026-04-18 | Progetto: radiOOracle / SportsOracle_Ultra
# Tutti i diritti riservati — All rights reserved.
"""
build_kdp_docx.py — Genera DOCX per Amazon KDP dai file Markdown aggiornati.
Converte oracle_factor_book_EN.md e oracle_factor_book_IT.md in DOCX
con formattazione KDP-ready: font Times 12pt, paragrafi giustificati,
titoli h1/h2/h3 correttamente stilati, tabelle, codice inline.
"""
import os
import sys, re
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installa python-docx: pip install python-docx")
    sys.exit(1)

BOOK_DIR = Path("os.getenv("USER_HOME", str(Path.home())) + "/Desktop/GUMROAD_PACKAGES/OracleBook")

SOURCES = [
    ("oracle_factor_book_IT.md", "IL_FATTORE_ORACLE_KDP_v2.docx"),
    ("oracle_factor_book_EN.md", "THE_ORACLE_FACTOR_KDP_v2.docx"),
]


def set_paragraph_justified(para):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    """Configura stili base KDP."""
    # Normal
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(18)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def parse_inline(run_container, text):
    """Aggiunge testo con **bold**, *italic*, `code` inline."""
    # Tokenizza bold, italic, code
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        # testo normale prima
        if m.start() > last:
            run = run_container.add_run(text[last:m.start()])
        if m.group(0).startswith("**"):
            run = run_container.add_run(m.group(2))
            run.bold = True
        elif m.group(0).startswith("*"):
            run = run_container.add_run(m.group(3))
            run.italic = True
        elif m.group(0).startswith("`"):
            run = run_container.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(11)
        last = m.end()
    if last < len(text):
        run_container.add_run(text[last:])


def add_table_from_md(doc, lines):
    """Crea tabella da righe markdown (| col | col |)."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-: ]+\|', line):
            continue  # separatore
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < ncols:
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                parse_inline(p, cell_text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()  # spazio dopo tabella


def add_code_block(doc, lines):
    """Aggiunge blocco di codice con font monospazio."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    code_text = "\n".join(lines)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def process_markdown(doc, md_text):
    """Processa il markdown riga per riga e popola il documento."""
    lines = md_text.splitlines()
    i = 0
    table_buf = []
    code_buf = []
    in_code = False
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Blocco di codice
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                add_code_block(doc, code_buf)
                code_buf = []
            else:
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Tabella
        if re.match(r'^\s*\|', line):
            table_buf.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            add_table_from_md(doc, table_buf)
            table_buf = []
            in_table = False
            # non incrementa, riprocessa riga corrente

        # Separatore ---
        if re.match(r'^-{3,}\s*$', line):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Heading 1 (#)
        m = re.match(r'^# (.+)', line)
        if m:
            p = doc.add_heading(m.group(1), level=1)
            i += 1
            continue

        # Heading 2 (##)
        m = re.match(r'^## (.+)', line)
        if m:
            p = doc.add_heading(m.group(1), level=2)
            i += 1
            continue

        # Heading 3 (###)
        m = re.match(r'^### (.+)', line)
        if m:
            p = doc.add_heading(m.group(1), level=3)
            i += 1
            continue

        # Lista puntata
        m = re.match(r'^(\s*)[*\-] (.+)', line)
        if m:
            indent = len(m.group(1)) // 2
            para = doc.add_paragraph(style="List Bullet")
            parse_inline(para, m.group(2))
            set_paragraph_justified(para)
            i += 1
            continue

        # Lista numerata
        m = re.match(r'^\s*\d+\. (.+)', line)
        if m:
            para = doc.add_paragraph(style="List Number")
            parse_inline(para, m.group(1))
            set_paragraph_justified(para)
            i += 1
            continue

        # Riga vuota
        if not line.strip():
            i += 1
            continue

        # Paragrafo normale
        para = doc.add_paragraph()
        set_paragraph_justified(para)
        parse_inline(para, line.strip())
        i += 1

    # Flush tabella aperta
    if in_table and table_buf:
        add_table_from_md(doc, table_buf)


def build_docx(md_path: Path, out_path: Path):
    print(f"Elaboro: {md_path.name} → {out_path.name}")
    md_text = md_path.read_text(encoding="utf-8")

    doc = Document()
    configure_styles(doc)

    # Margini KDP standard (2.54cm = 1 pollice)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    process_markdown(doc, md_text)

    doc.save(out_path)
    size_kb = out_path.stat().st_size // 1024
    print(f"  Salvato: {out_path} ({size_kb} KB)")
    return out_path


if __name__ == "__main__":
    for md_name, docx_name in SOURCES:
        md_path   = BOOK_DIR / md_name
        docx_path = BOOK_DIR / docx_name
        if not md_path.exists():
            print(f"[SKIP] {md_name} non trovato")
            continue
        build_docx(md_path, docx_path)
    print("\nFatto.")