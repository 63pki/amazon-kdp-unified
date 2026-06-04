#!/usr/bin/env python3
"""
Kindle Manuscript Compiler
Converts a book project's chapters into a formatted .docx file
with proper front matter, back matter, and Kindle-specific formatting.

Supports two modes:
  1. chapters/ directory — assembles from individual ch01.md..chNN.md files
  2. kindle_manuscript.md — legacy single-file mode (backward compatible)

Usage:
    python compile_kindle.py output/my-book                            # Auto-detect mode
    python compile_kindle.py output/my-book --metadata custom.json     # Custom metadata
    python compile_kindle.py                                           # Legacy: hardcoded paths
"""

import argparse
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from diagram_renderer import parse_diagram_spec, render_diagram, extract_diagram_blocks
    HAS_DIAGRAM_RENDERER = True
except ImportError:
    HAS_DIAGRAM_RENDERER = False


# ── Legacy Defaults (backward compatibility) ──────────────────────────────

_LEGACY_DEFAULTS = {
    "book_title": "Is AI helpful for PM, or is PM helpful for AI?",
    "subtitle": "A Software Engineer's Guide to Upskilling Product Managers in the Vibe Coding Era",
    "author": "Herschel J.",
    "publisher": "Independently Published",
    "copyright_year": "2026",
    "dedication": (
        "For the product managers who asked the right questions at the right time\u2014"
        "and the engineers who learned to teach instead of gatekeep."
    ),
    "about_author": (
        "is a software engineer who got tired of being the bottleneck. After years of "
        "building mockups at 10 PM because PMs couldn't build their own, he started "
        "teaching them. Prompt engineering. Vibe coding. How to talk to AI tools like "
        "a colleague instead of a magic eight ball. Within months, his PMs were "
        "prototyping features, automating their own busywork, and shipping demos "
        "without touching an engineer. This book is what he taught them."
    ),
    "also_by": [
        "The AI-Fluent PM Workbook (Coming Soon)",
        "The PM's Prompt Library: 50 AI Prompts for Product Work (Coming Soon)",
    ],
}

# Legacy markers for single-file mode
_BODY_START_MARKER = "# PART 1: THE QUESTION"
_BODY_END_MARKER = "# About the Author"


# ── Metadata Loading ──────────────────────────────────────────────────────


def load_metadata(project_dir: Path, metadata_path: Path | None = None) -> dict:
    """Load book metadata from metadata.json or fall back to legacy defaults."""
    # Try explicit metadata path first
    if metadata_path and metadata_path.exists():
        with open(metadata_path, encoding="utf-8") as f:
            meta = json.load(f)
        print(f"  Loaded metadata from {metadata_path}")
        return {**_LEGACY_DEFAULTS, **meta}

    # Try project_dir/metadata.json
    default_meta = project_dir / "metadata.json"
    if default_meta.exists():
        with open(default_meta, encoding="utf-8") as f:
            meta = json.load(f)
        print(f"  Loaded metadata from {default_meta}")
        merged = {**_LEGACY_DEFAULTS, **meta}
        # Ensure also_by is always a list
        if not isinstance(merged.get("also_by"), list):
            merged["also_by"] = [str(merged["also_by"])] if merged.get("also_by") else []
        return merged

    print("  No metadata.json found — using legacy defaults")
    return dict(_LEGACY_DEFAULTS)


# ── Chapter Assembly ──────────────────────────────────────────────────────


def parse_part_structure(project_dir: Path) -> list[dict]:
    """
    Parse the ghostwriter output to determine PART structure.
    Returns a list of {"title": "PART N: ...", "chapters": [1, 2, 3]} dicts.
    """
    ghostwriter_path = project_dir / "ghostwriter.md"
    if not ghostwriter_path.exists():
        return []

    text = ghostwriter_path.read_text(encoding="utf-8")
    parts = []
    current_part = None

    for line in text.split("\n"):
        # Match ### PART N: TITLE (possibly with parens at end)
        part_match = re.match(r"^### (PART \d+:\s*.+?)(?:\s*\(.*\))?\s*$", line)
        if part_match:
            current_part = {"title": part_match.group(1).strip(), "chapters": []}
            parts.append(current_part)
            continue

        # Match **Chapter N: Title**
        ch_match = re.match(r"^\*\*Chapter (\d+):", line)
        if ch_match and current_part is not None:
            current_part["chapters"].append(int(ch_match.group(1)))

    return parts


def assemble_from_chapters(project_dir: Path) -> str:
    """
    Assemble the full manuscript body from individual chapter files in chapters/.
    Inserts PART headings between chapter groups based on ghostwriter outline.
    """
    chapters_dir = project_dir / "chapters"
    if not chapters_dir.exists():
        return ""

    # Find all chapter files
    chapter_files = sorted(chapters_dir.glob("ch*.md"))
    if not chapter_files:
        return ""

    # Parse which chapters belong to which parts
    parts = parse_part_structure(project_dir)

    # Build a map: chapter_num -> part_title (for the FIRST chapter in each part)
    part_starts = {}
    for part in parts:
        if part["chapters"]:
            part_starts[part["chapters"][0]] = part["title"]

    # Assemble
    body_parts = []
    for ch_file in chapter_files:
        # Extract chapter number from filename (ch01.md -> 1)
        num_match = re.match(r"ch(\d+)\.md", ch_file.name)
        if not num_match:
            continue
        ch_num = int(num_match.group(1))

        # Insert PART heading if this chapter starts a new part
        if ch_num in part_starts:
            body_parts.append(f"# {part_starts[ch_num]}\n")

        # Read chapter content
        chapter_text = ch_file.read_text(encoding="utf-8").strip()
        body_parts.append(chapter_text)

    return "\n\n".join(body_parts)


# ── Page Numbers ──────────────────────────────────────────────────────────


def add_page_numbers(doc):
    """Add centered page numbers to document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add PAGE field
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

        # Style the page number
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.name = 'Georgia'
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Chapter Visuals ──────────────────────────────────────────────────────


def get_chapter_visuals(project_dir: Path, chapter_num: int) -> list[Path]:
    """Find visual PNG files for a given chapter number."""
    visuals_dir = project_dir / "chapters" / "visuals"
    if not visuals_dir.exists():
        return []
    pattern = f"ch{chapter_num:02d}_*.png"
    return sorted(visuals_dir.glob(pattern))


# ── Styles ────────────────────────────────────────────────────────────────


def setup_styles(doc):
    """Configure document styles for Kindle formatting."""
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.first_line_indent = Inches(0.3)
    pf.line_spacing = 1.5

    # No-indent paragraph (after headings and breaks)
    no_indent = doc.styles.add_style('NoIndent', WD_STYLE_TYPE.PARAGRAPH)
    no_indent.base_style = doc.styles['Normal']
    no_indent.paragraph_format.first_line_indent = Inches(0)

    # Heading 1 — Chapter/Part titles (Kindle TOC generation)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Helvetica Neue'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)
    h1.paragraph_format.space_before = Pt(48)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.page_break_before = True
    # Add decorative bottom border to H1
    pPr = h1.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), '1A2B4A')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Heading 2 — Section headings
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Helvetica Neue'
    h2.font.size = Pt(17)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h2.paragraph_format.space_before = Pt(28)
    h2.paragraph_format.space_after = Pt(10)
    h2.paragraph_format.first_line_indent = Inches(0)

    # Heading 3 — Subsection headings
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Helvetica Neue'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2A, 0x4F, 0x6F)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Inches(0)

    # Heading 4 — Sub-subsection headings
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Helvetica Neue'
    h4.font.size = Pt(11.5)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x4A, 0x5A, 0x7A)
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(4)
    h4.paragraph_format.first_line_indent = Inches(0)

    # Section break style
    section_break = doc.styles.add_style('SectionBreak', WD_STYLE_TYPE.PARAGRAPH)
    section_break.font.name = 'Georgia'
    section_break.font.size = Pt(14)
    section_break.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section_break.paragraph_format.space_before = Pt(18)
    section_break.paragraph_format.space_after = Pt(18)
    section_break.paragraph_format.first_line_indent = Inches(0)

    # Centered style (for title pages etc.)
    centered = doc.styles.add_style('Centered', WD_STYLE_TYPE.PARAGRAPH)
    centered.base_style = doc.styles['Normal']
    centered.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered.paragraph_format.first_line_indent = Inches(0)

    # Code block style
    code = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = 'Courier New'
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.first_line_indent = Inches(0)
    code.paragraph_format.left_indent = Inches(0.3)

    # List item style
    list_item = doc.styles.add_style('ListItem', WD_STYLE_TYPE.PARAGRAPH)
    list_item.base_style = doc.styles['Normal']
    list_item.paragraph_format.first_line_indent = Inches(0)
    list_item.paragraph_format.left_indent = Inches(0.3)
    list_item.paragraph_format.space_before = Pt(2)
    list_item.paragraph_format.space_after = Pt(2)

    # Table cell style
    table_cell = doc.styles.add_style('TableCell', WD_STYLE_TYPE.PARAGRAPH)
    table_cell.font.name = 'Georgia'
    table_cell.font.size = Pt(9.5)
    table_cell.paragraph_format.space_before = Pt(2)
    table_cell.paragraph_format.space_after = Pt(2)
    table_cell.paragraph_format.first_line_indent = Inches(0)

    return doc


# ── Front & Back Matter ───────────────────────────────────────────────────


def add_front_matter(doc, meta: dict):
    """Add Kindle front matter: title page, copyright, TOC placeholder."""
    book_title = meta["book_title"]
    subtitle = meta.get("subtitle", "")
    author = meta["author"]
    publisher = meta.get("publisher", "Independently Published")
    copyright_year = meta.get("copyright_year", "2026")
    dedication = meta.get("dedication", "")

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph('', style='Centered')

    p = doc.add_paragraph(style='Centered')
    run = p.add_run(book_title)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4A)

    doc.add_paragraph('', style='Centered')

    if subtitle:
        p = doc.add_paragraph(style='Centered')
        run = p.add_run(subtitle)
        run.font.name = 'Helvetica Neue'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x7A)

    for _ in range(4):
        doc.add_paragraph('', style='Centered')

    p = doc.add_paragraph(style='Centered')
    run = p.add_run(author)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(16)
    run.font.bold = True

    # ── Copyright Page ──
    doc.add_page_break()

    title_line = f'{book_title}: {subtitle}' if subtitle else book_title
    p = doc.add_paragraph(style='NoIndent')
    run = p.add_run(title_line)
    run.font.size = Pt(10)
    run.font.bold = True

    doc.add_paragraph('', style='NoIndent')

    copyright_text = f"Copyright \u00A9 {copyright_year} {author}. All rights reserved."
    p = doc.add_paragraph(copyright_text, style='NoIndent')
    p.runs[0].font.size = Pt(9.5)

    doc.add_paragraph('', style='NoIndent')

    legal_text = (
        "No part of this book may be reproduced in any form without written "
        "permission from the publisher, except for brief quotations in reviews."
    )
    p = doc.add_paragraph(legal_text, style='NoIndent')
    p.runs[0].font.size = Pt(9.5)

    doc.add_paragraph('', style='NoIndent')

    p = doc.add_paragraph(f"Published by {publisher}", style='NoIndent')
    p.runs[0].font.size = Pt(9.5)

    doc.add_paragraph('', style='NoIndent')

    p = doc.add_paragraph(
        "The information in this book is provided for educational and informational "
        "purposes only. While every effort has been made to ensure accuracy, the author "
        "and publisher make no warranties regarding the completeness or accuracy of the "
        "contents. AI tools, platforms, and capabilities described in this book may change "
        "after publication.",
        style='NoIndent'
    )
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.italic = True

    # ── Dedication Page ──
    if dedication:
        doc.add_page_break()
        for _ in range(8):
            doc.add_paragraph('', style='Centered')

        p = doc.add_paragraph(style='Centered')
        run = p.add_run(dedication)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.font.italic = True

    # ── Table of Contents placeholder ──
    doc.add_page_break()
    p = doc.add_paragraph('Table of Contents', style='Heading 1')
    p.paragraph_format.page_break_before = False

    p = doc.add_paragraph(
        'This table of contents is automatically generated by your Kindle device.',
        style='NoIndent'
    )
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)


def add_back_matter(doc, meta: dict):
    """Add Kindle back matter: about author, also by, review request."""
    author = meta["author"]
    about_author = meta.get("about_author", "")
    also_by = meta.get("also_by", [])

    # ── About the Author ──
    if about_author:
        doc.add_paragraph('About the Author', style='Heading 1')
        p = doc.add_paragraph(style='NoIndent')
        run = p.add_run(f'{author} ')
        run.font.bold = True
        p.add_run(about_author)

    # ── Also By ──
    if also_by:
        doc.add_paragraph(f'Also By {author}', style='Heading 1')
        for title in also_by:
            p = doc.add_paragraph(style='Centered')
            run = p.add_run(title)
            run.font.italic = True

    # ── Review Request ──
    doc.add_paragraph('One Last Thing\u2026', style='Heading 1')

    doc.add_paragraph(
        'If you enjoyed this book, please consider leaving a review on Amazon. '
        'Your feedback helps independent authors and helps other readers discover '
        'this book.',
        style='NoIndent'
    )

    doc.add_paragraph(
        'Even two or three sentences make a real difference. Thank you for reading.',
        style='Normal'
    )


# ── Markdown → Docx Conversion ────────────────────────────────────────────


def add_inline_formatting(paragraph, text):
    """Parse markdown inline formatting (bold, italic, code) and add runs."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """Add a professionally formatted table to the document."""
    if not rows or not rows[0]:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.style = doc.styles['TableCell']
                add_inline_formatting(p, cell_text)

                # Header row styling
                if i == 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '1A2B4A')
                    shading.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Alternating row shading
                elif i % 2 == 1:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F5F5F5')
                    shading.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph('', style='NoIndent')


def convert_markdown_to_docx(doc, markdown_text, project_dir=None):
    """Convert manuscript markdown to docx paragraphs."""
    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    after_heading = True
    current_chapter = 0
    chapter_first_break_done = set()  # Track which chapters have had their visual inserted

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Track chapter numbers
        ch_heading = re.match(r'^# Chapter (\d+):', stripped)
        if ch_heading:
            current_chapter = int(ch_heading.group(1))

        # ── Diagram blocks ──
        if stripped == '```diagram' and HAS_DIAGRAM_RENDERER and project_dir:
            # Collect the diagram spec lines until closing ```
            i += 1
            diagram_lines = []
            while i < len(lines) and lines[i].strip() != '```':
                diagram_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            # Parse and render
            spec = parse_diagram_spec('\n'.join(diagram_lines))
            if spec.get('type'):
                vis_dir = project_dir / 'chapters' / 'visuals' / 'inline'
                img_path = render_diagram(spec, vis_dir)
                if img_path and img_path.exists():
                    doc.add_paragraph('', style='NoIndent')
                    p = doc.add_paragraph(style='Centered')
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(4.5))
                    # Add caption if title exists
                    if spec.get('title'):
                        cap = doc.add_paragraph(style='Centered')
                        cap.paragraph_format.space_before = Pt(2)
                        cap.paragraph_format.space_after = Pt(12)
                        cap_run = cap.add_run(spec['title'])
                        cap_run.font.size = Pt(8.5)
                        cap_run.font.italic = True
                        cap_run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x7A)
                    else:
                        doc.add_paragraph('', style='NoIndent')
            after_heading = True
            continue

        # ── Code blocks ──
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph('', style='NoIndent')
                after_heading = True
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            doc.add_paragraph(line.rstrip(), style='CodeBlock')
            i += 1
            continue

        # ── Empty lines ──
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule / section break ──
        if stripped in ('---', '***', '* * *'):
            doc.add_paragraph('*  *  *', style='SectionBreak')
            after_heading = True

            # Insert visuals after first section break in each chapter
            if project_dir and current_chapter > 0 and current_chapter not in chapter_first_break_done:
                chapter_first_break_done.add(current_chapter)
                visuals = get_chapter_visuals(project_dir, current_chapter)
                for vis_path in visuals:
                    doc.add_paragraph('', style='NoIndent')
                    p = doc.add_paragraph(style='Centered')
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    run = p.add_run()
                    run.add_picture(str(vis_path), width=Inches(4.5))
                    doc.add_paragraph('', style='NoIndent')

            i += 1
            continue

        # ── Tables ──
        if stripped.startswith('|'):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            after_heading = True
            continue

        # ── Headings ──
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
            doc.add_paragraph(text, style=style_map.get(level, 'Heading 4'))
            after_heading = True
            i += 1
            continue

        # ── Bullet lists ──
        if re.match(r'^[-*]\s', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='ListItem')
            p.add_run('\u2022  ')
            add_inline_formatting(p, text)
            after_heading = False
            i += 1
            continue

        # ── Numbered lists ──
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph(style='ListItem')
            run = p.add_run(f'{num}.  ')
            run.font.bold = True
            add_inline_formatting(p, text)
            after_heading = False
            i += 1
            continue

        # ── Indented list items (sub-bullets) ──
        if re.match(r'^\s+[-*]\s', line):
            text = re.sub(r'^\s+[-*]\s+', '', line)
            p = doc.add_paragraph(style='ListItem')
            p.paragraph_format.left_indent = Inches(0.6)
            p.add_run('\u25E6  ')
            add_inline_formatting(p, text)
            i += 1
            continue

        # ── Regular paragraphs ──
        style = 'NoIndent' if after_heading else 'Normal'
        p = doc.add_paragraph(style=style)
        add_inline_formatting(p, stripped)
        after_heading = False
        i += 1


# ── Body Extraction ───────────────────────────────────────────────────────


def extract_body_from_manuscript(markdown_text: str) -> str:
    """Extract book body from a single kindle_manuscript.md file (legacy mode)."""
    start_idx = markdown_text.find(_BODY_START_MARKER)
    end_idx = markdown_text.find(_BODY_END_MARKER)

    if start_idx == -1:
        raise ValueError(f"Could not find body start marker: {_BODY_START_MARKER}")

    body = markdown_text[start_idx:end_idx] if end_idx != -1 else markdown_text[start_idx:]
    return body.strip()


def count_words(text: str) -> int:
    """Count words in text."""
    return len(text.split())


# ── Main Compilation ──────────────────────────────────────────────────────


def compile_book(project_dir: Path, metadata_path: Path | None = None) -> dict:
    """
    Compile a book project into a Kindle .docx manuscript.

    Tries chapters/ directory first, then falls back to kindle_manuscript.md.

    Returns a summary dict with word count, chapter count, etc.
    """
    print("=" * 60)
    print("KINDLE MANUSCRIPT COMPILER")
    print("=" * 60)

    # Load metadata
    meta = load_metadata(project_dir, metadata_path)
    print(f"\n  Title:  {meta['book_title']}")
    print(f"  Author: {meta['author']}")

    # Determine source mode: chapters/ directory or single file
    chapters_dir = project_dir / "chapters"
    manuscript_path = project_dir / "kindle_manuscript.md"

    if chapters_dir.exists() and list(chapters_dir.glob("ch*.md")):
        # Chapters mode
        print(f"\nAssembling from {chapters_dir}/...")
        body = assemble_from_chapters(project_dir)
        if not body:
            raise ValueError(f"No chapter files found in {chapters_dir}")
    elif manuscript_path.exists():
        # Legacy single-file mode
        print(f"\nReading manuscript from {manuscript_path}...")
        raw = manuscript_path.read_text(encoding="utf-8")
        body = extract_body_from_manuscript(raw)
    else:
        raise FileNotFoundError(
            f"No chapters/ directory or kindle_manuscript.md found in {project_dir}"
        )

    word_count = count_words(body)
    chapter_count = len(re.findall(r'^# Chapter \d+:', body, re.MULTILINE))
    part_count = len(re.findall(r'^# PART \d+:', body, re.MULTILINE))

    print(f"  Body word count: {word_count:,}")
    print(f"  Chapters: {chapter_count}")
    print(f"  Parts: {part_count}")

    # Create document
    print("\nCreating .docx document...")
    doc = Document()

    # Set up page margins for Kindle
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.different_first_page_header_footer = False

    doc = setup_styles(doc)
    add_page_numbers(doc)

    print("  Adding front matter...")
    add_front_matter(doc, meta)

    print("  Converting manuscript body...")
    convert_markdown_to_docx(doc, body, project_dir=project_dir)

    print("  Adding back matter...")
    add_back_matter(doc, meta)

    # Save
    output_path = project_dir / "kindle_manuscript.docx"
    print(f"\nSaving to {output_path}...")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    file_size = output_path.stat().st_size
    total_paras = len(doc.paragraphs)
    total_tables = len(doc.tables)

    print(f"  File size: {file_size:,} bytes ({file_size / 1024:.0f} KB)")

    print("\n" + "=" * 60)
    print("COMPILATION COMPLETE")
    print("=" * 60)
    print(f"  Title: {meta['book_title']}")
    print(f"  Author: {meta['author']}")
    print(f"  Word count: {word_count:,}")
    print(f"  Chapters: {chapter_count}")
    print(f"  Parts: {part_count}")
    print(f"  Paragraphs: {total_paras}")
    print(f"  Tables: {total_tables}")
    print(f"  Output: {output_path}")

    summary = {
        'word_count': word_count,
        'chapter_count': chapter_count,
        'part_count': part_count,
        'total_paras': total_paras,
        'total_tables': total_tables,
        'file_size': file_size,
    }
    return summary


def main():
    parser = argparse.ArgumentParser(
        description="Compile a book project into a Kindle .docx manuscript.",
    )
    parser.add_argument(
        "project_dir",
        type=Path,
        nargs="?",
        default=None,
        help="Path to the project output directory (e.g. output/my-book)",
    )
    parser.add_argument(
        "--metadata",
        type=Path,
        default=None,
        help="Path to a custom metadata.json file",
    )

    args = parser.parse_args()

    if args.project_dir is None:
        # Legacy mode: use hardcoded ai-pm-era path
        args.project_dir = Path("output/ai-pm-era")

    compile_book(args.project_dir, args.metadata)


if __name__ == '__main__':
    main()